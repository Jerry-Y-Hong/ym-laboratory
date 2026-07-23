import pandas as pd
import os
import docx

root_dir = r"C:\Users\car13\OneDrive\MFCO_MASTER_RECOVERY"
baseline_dir = os.path.join(root_dir, "01_BASELINE")
out_file = os.path.join(root_dir, "scratch", "baseline_inspection_report.txt")

with open(out_file, "w", encoding="utf-8") as f:
    f.write("=== BASELINE DIRECTORY DETAILED INSPECTION ===\n\n")
    
    # 1. Inspect 식품성분표(10개정판).xlsx
    food_comp_path = os.path.join(baseline_dir, "식품성분표(10개정판).xlsx")
    if os.path.exists(food_comp_path):
        f.write(f"1. 식품성분표(10개정판).xlsx Found.\n")
        try:
            xl = pd.ExcelFile(food_comp_path)
            f.write(f"Sheets: {xl.sheet_names}\n")
            for sheet in xl.sheet_names[:3]:  # Check first few sheets
                df = pd.read_excel(food_comp_path, sheet_name=sheet, nrows=5)
                f.write(f"Sheet: {sheet}, Shape: {df.shape} (preview of first 5 rows)\n")
                f.write(f"Columns: {list(df.columns)}\n")
                f.write(df.head(2).to_string())
                f.write("\n\n")
        except Exception as e:
            f.write(f"Error reading 식품성분표: {e}\n\n")
    else:
        f.write("식품성분표(10개정판).xlsx Not Found.\n\n")
        
    # 2. Inspect MFCO_PHASE1_MASTER_CONSTITUTION_v1.0.docx
    doc_path = os.path.join(baseline_dir, "MFCO_PHASE1_MASTER_CONSTITUTION_v1.0.docx")
    if os.path.exists(doc_path):
        f.write("2. MFCO_PHASE1_MASTER_CONSTITUTION_v1.0.docx Found.\n")
        try:
            doc = docx.Document(doc_path)
            f.write(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}\n\n")
            f.write("=== Title / First 20 Paragraphs ===\n")
            for idx, p in enumerate(doc.paragraphs[:30]):
                f.write(f"P{idx}: {p.text}\n")
            
            f.write("\n=== First Table Content (up to 5 rows) ===\n")
            if len(doc.tables) > 0:
                table = doc.tables[0]
                for r_idx, row in enumerate(table.rows[:10]):
                    cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    f.write(f"Row{r_idx}: {cells}\n")
            else:
                f.write("No tables in docx\n")
        except Exception as e:
            f.write(f"Error reading docx: {e}\n\n")
    else:
        f.write("MFCO_PHASE1_MASTER_CONSTITUTION_v1.0.docx Not Found.\n\n")

print("Written inspection report to:", out_file)
