import docx
import os

root_dir = r"C:\Users\car13\OneDrive\MFCO_MASTER_RECOVERY"
docx_path = os.path.join(root_dir, "06_DOCUMENT", "한약재_효능_통합정리_보고서.docx")
out_path = os.path.join(root_dir, "scratch", "docx_content.txt")

with open(out_path, "w", encoding="utf-8") as f:
    if os.path.exists(docx_path):
        try:
            doc = docx.Document(docx_path)
            f.write(f"Number of paragraphs: {len(doc.paragraphs)}\n")
            f.write(f"Number of tables: {len(doc.tables)}\n\n")
            
            f.write("=== First 20 paragraphs ===\n")
            for i, p in enumerate(doc.paragraphs[:20]):
                f.write(f"P{i}: {p.text}\n")
                
            f.write("\n=== First Table Content (first 5 rows) ===\n")
            if len(doc.tables) > 0:
                table = doc.tables[0]
                for r_idx, row in enumerate(table.rows[:10]):
                    cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    f.write(f"Row{r_idx}: {cells}\n")
            else:
                f.write("No tables found in docx\n")
        except Exception as e:
            f.write(f"Error reading docx: {e}\n")
    else:
        f.write("Docx file not found\n")

print("Docx inspection written to:", out_path)
